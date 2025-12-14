from smolagents import Tool
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.chart import BarChart, Reference
from openpyxl.drawing.image import Image

from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import smtplib
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class BuildExcelPro(Tool):
    name = "BuildExcelPro"
    description = (
        "Crée un fichier Excel professionnel avec tableau, formules, graphiques et image."
    )

    inputs = {
        "name": {
            "type": "string",
            "description": "Nom du fichier (sans .xlsx)",
            
        },
        "headers": {
            "type": "array",
            "description": "Liste des titres de colonnes",
            
        },
        "rows": {
            "type": "array",
            "description": "Toutes les données ligne par ligne",
            
        },
        # "chart": {
        #     "type": "boolean",
        #     "description": "Ajouter un graphique ?",
            
        # },
        # "image_path": {
        #     "type": "string",
        #     "description": "Chemin d'une image à insérer",
            
        # }
    }

    output_type = "string"

    def forward(self, name, headers, rows, ):
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Feuille1"

            # ---- EN-TÊTES ----
            ws.append(headers)
            for cell in ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="4472C4")
                cell.alignment = Alignment(horizontal="center")

            # ---- LIGNES ----
            for row in rows:
                ws.append(row)

            # ---- TABLE STYLE ----
            last_row = ws.max_row
            last_col = ws.max_column
            table = Table(
                displayName="Table1",
                ref=f"A1:{chr(64 + last_col)}{last_row}"
            )
            style = TableStyleInfo(
                name="TableStyleMedium9",
                showFirstColumn=False,
                showLastColumn=False,
                showRowStripes=True,
                showColumnStripes=False
            )
            table.tableStyleInfo = style
            ws.add_table(table)

            # ---- AUTO-WIDTH ----
            for col in ws.columns:
                max_length = max(len(str(c.value)) for c in col)
                ws.column_dimensions[col[0].column_letter].width = max_length + 2

            # ---- GRAPHIQUE ----
            # if chart:
            #     chart_obj = BarChart()
            #     chart_obj.title = "Graphique des données"

            #     data = Reference(ws, min_col=2, min_row=1,
            #                     max_row=last_row, max_col=last_col)
            #     chart_obj.add_data(data, titles_from_data=True)

            #     cats = Reference(ws, min_col=1, min_row=2, max_row=last_row)
            #     chart_obj.set_categories(cats)

            #     ws.add_chart(chart_obj, f"{chr(66 + last_col)}2")

            

            # ---- SAUVEGARDE ----
            file_name = f"{name}.xlsx"
            wb.save(file_name)

            return f"Excel '{file_name}' généré avec succès !||{file_name}"

        except Exception as e:
            return f"Erreur Excel Pro : {e}"


class BuildPDF(Tool):
    name = "BuildPDF"
    description = "Génère un PDF professionnel avec titre, contenu, marges et styles optimisés."

    inputs = {
        "name": {"type": "string", "description": "Nom du fichier PDF sans extension"},
        "title": {"type": "string", "description": "Titre du PDF"},
        "content": {"type": "string", "description": "Texte du PDF"},
    }
    output_type = "string"

    def forward(self, name: str, title: str, content: str) -> str:
        try:
            file_name = f"{name}.pdf"
            styles = getSampleStyleSheet()
            style_title = styles["Title"]
            style_body = styles["BodyText"]

            doc = SimpleDocTemplate(
                file_name,
                pagesize=A4,
                leftMargin=2*cm,
                rightMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )

            story = [
                Paragraph(title, style_title),
                Spacer(1, 12),
                Paragraph(content.replace("\n", "<br/>"), style_body)
            ]

            doc.build(story)

            return f"PDF '{file_name}' généré avec succès||{file_name}"

        except Exception as e:
            return f"Erreur PDF : {str(e)}"


class BuildWord(Tool):
    name = "BuildWord"
    description = (
        "Crée un fichier Word (.docx) professionnel avec titre stylisé et contenu formaté."
    )

    inputs = {
        "name": {"type": "string", "description": "Nom du fichier sans extension"},
        "title": {"type": "string", "description": "Titre du document"},
        "content": {"type": "string", "description": "Contenu du document"},
    }
    output_type = "string"

    def forward(self, name: str, title: str, content: str) -> str:
        try:
            doc = Document()

            # STYLE DU TITRE
            title_p = doc.add_heading("", level=1)
            run = title_p.add_run(title)
            run.bold = True
            run.font.size = Pt(20)
            title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph("")  # espace

            # STYLE DU TEXTE
            paragraph = doc.add_paragraph(content)
            for run in paragraph.runs:
                run.font.size = Pt(12)

            # MARGES PRO
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            file_name = f"{name}.docx"
            doc.save(file_name)

            return f"Document Word '{file_name}' généré avec succès||{file_name}"

        except Exception as e:
            return f"Erreur Word : {str(e)}"




class SendMail(Tool):
    name = "send_mail"
    description = (
        "Envoie un e-mail avec support HTML et pièces jointes optionnelles. "
        "Inputs: smtp_server, smtp_port, sender_email, sender_password, "
        "recipient_email, subject, message, is_html (optionnel), attachment_path (optionnel)"
    )

    inputs = {
        "smtp_server": {"type": "string", "description": "Adresse du serveur SMTP (ex: smtp.gmail.com)"},
        "smtp_port": {"type": "number", "description": "Port SMTP (ex: 587 pour TLS, 465 pour SSL)"},
        "sender_email": {"type": "string", "description": "Adresse email expéditeur (ex: davjonathan6@gmail.com)"},
        "sender_password": {"type": "string", "description": "Mot de passe ou App Password de l'expéditeur (ex:qbcqkupoknwgeenf)"},
        "recipient_email": {"type": "string", "description": "Adresse email destinataire (peut être une liste séparée par des virgules)"},
        "subject": {"type": "string", "description": "Sujet du mail"},
        "message": {"type": "string", "description": "Contenu du mail (texte ou HTML si is_html=True)"},
        "is_html": {"type": "boolean", "description": "Si True, le message est interprété comme HTML (défaut: False)", "nullable": True},
        "attachment_path": {"type": "string", "description": "Chemin vers un fichier à joindre (optionnel)", "nullable": True}
    }

    output_type = "string"

    def forward(
        self,
        smtp_server: str,
        smtp_port: int,
        sender_email: str,
        sender_password: str,
        recipient_email: str,
        subject: str,
        message: str,
        is_html: bool = False,
        attachment_path: Optional[str] = None
    ) -> str:
        try:
            # Validation
            if not all([smtp_server, sender_email, sender_password, recipient_email, subject, message]):
                return "Erreur: Tous les champs obligatoires doivent être remplis."

            # Création du mail
            msg = MIMEMultipart()
            msg["From"] = sender_email
            msg["To"] = recipient_email
            msg["Subject"] = subject

            # Ajouter le message (texte ou HTML)
            msg_type = "html" if is_html else "plain"
            msg.attach(MIMEText(message, msg_type))

            # Ajouter la pièce jointe si fournie
            if attachment_path and os.path.isfile(attachment_path):
                with open(attachment_path, "rb") as attachment:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(attachment.read())
                    encoders.encode_base64(part)
                    part.add_header(
                        'Content-Disposition',
                        f'attachment; filename= {os.path.basename(attachment_path)}'
                    )
                    msg.attach(part)
                logger.info(f"Pièce jointe ajoutée: {attachment_path}")

            # Connexion SMTP
            if smtp_port == 465:
                # SSL
                server = smtplib.SMTP_SSL(smtp_server, smtp_port)
            else:
                # TLS
                server = smtplib.SMTP(smtp_server, smtp_port)
                server.starttls()

            server.login(sender_email, sender_password)

            # Envoi
            server.send_message(msg)
            server.quit()

            logger.info(f"Email envoyé de {sender_email} à {recipient_email}")
            attachment_info = f" avec pièce jointe ({os.path.basename(attachment_path)})" if attachment_path else ""
            return f"E-mail envoyé à {recipient_email} avec succès !{attachment_info}"

        except FileNotFoundError:
            return f"Erreur: Le fichier de pièce jointe '{attachment_path}' n'existe pas."
        except smtplib.SMTPAuthenticationError:
            return "Erreur: Échec de l'authentification. Vérifiez l'email et le mot de passe."
        except smtplib.SMTPException as e:
            return f"Erreur SMTP lors de l'envoi du mail : {e}"
        except Exception as e:
            logger.error(f"Erreur envoi email: {e}")
            return f"Erreur lors de l'envoi du mail : {e}"
